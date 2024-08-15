import os
import tempfile
from pathlib import Path
import cv2
import streamlit as st
import torch
from accelerate import Accelerator
import torchvision.transforms as transforms
from PIL import Image
import io
import pytesseract as pyt
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

# Import Model and other necessary functions
from models import Model  # Assuming Model is the name of your model class
from utils.utils import load_checkpoint, seed_everything
from config import Config

pyt.pytesseract.tesseract_cmd = "C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
from docx.shared import Pt

def perform_ocr_and_generate_docx(uploaded_file):
    temp_image = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    temp_image.write(uploaded_file.read())
    temp_image.close()

    # Read the image using OpenCV
    img = cv2.imread(temp_image.name)

    # Perform OCR on the image and convert the output to plain text
    text = pyt.image_to_string(img)

    # Create a new Word document with the extracted text
    doc = Document()

    # Define a style with font size 9
    style = doc.styles.add_style('CustomStyle', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.size = Pt(9)

    # Add a paragraph with the defined style
    paragraph = doc.add_paragraph(style='CustomStyle')

    # Add the extracted text to the paragraph
    paragraph.add_run(text)

    # Remove the temporary image file
    os.unlink(temp_image.name)

    return doc, text

# Define sample_input function
def sample_input(image_path, model_path):
    accelerator = Accelerator()

    # Load the input image
    input_image = Image.open(image_path)
    opt = Config('config.yml')
    seed_everything(opt.OPTIM.SEED)

    # Apply the necessary transformations
    transform = transforms.Compose([
        transforms.ToTensor(),
    ])

    input_tensor = transform(input_image).unsqueeze(0).to(torch.device('cuda'))

    # Load the model
    model = Model()
    load_checkpoint(model, model_path)

    model, input_tensor = accelerator.prepare(model, input_tensor)
    model.eval()

    with torch.no_grad():
        output_tensor = model(input_tensor)
        output_image = transforms.ToPILImage()(output_tensor.squeeze().cpu())

    return output_image


def main():
    st.title("Shadow Removal App")
    st.write("Upload an image containing shadows to remove them.")

    # Path to your trained model file
    model_path = "./checkpoints/RDD_epoch_18.pth"

    uploaded_file = st.file_uploader("Choose an image...", type=["jpg", "jpeg", "png"])

    if uploaded_file is not None:
        # Display file metadata
        # st.write("Uploaded file metadata:", uploaded_file)

        # Display the uploaded image
        st.image(uploaded_file, caption="Uploaded Image", use_column_width=True)

        # Perform shadow removal
        output_image = sample_input(uploaded_file, model_path)

        # Display the shadow-free image
        st.image(output_image, caption="Shadow-Free Image", use_column_width=True)
        output_bytes = io.BytesIO()
        output_image.save(output_bytes, format='PNG')
        output_bytes.seek(0)
        uploaded_filename = uploaded_file.name
        # Extract the file extension
        file_ext = os.path.splitext(uploaded_filename)[1]
        # Generate the new filename with "_shfree" suffix
        download_filename = os.path.splitext(uploaded_filename)[0] + '_shfree' + file_ext

        st.download_button(label="Download Shadow-Free Image", data=output_bytes, file_name=download_filename,mime='image/png')
    st.title("OCR Document Processing")
    ocr_file = st.file_uploader("Upload Image for OCR", type=["jpg", "jpeg", "png"])
    if ocr_file is not None:
        # Display the uploaded image for OCR
        st.image(ocr_file, caption="Uploaded Image for OCR", use_column_width=True)

        if "extracted_text" not in st.session_state:
            st.session_state["extracted_text"] = ""

        if st.button("Perform OCR"):
            # Perform OCR processing and get the Word document and extracted text
            try:
                ocr_doc, extracted_text = perform_ocr_and_generate_docx(ocr_file)

                # Display the extracted text
                st.subheader("Extracted Text:")
                st.session_state["extracted_text"] = extracted_text
                st.text(extracted_text)

                # Create a new .docx file with the same name as the uploaded image and the suffix _ocr.docx
                doc_name = Path(ocr_file.name).with_suffix('.docx')
                ocr_doc.save(doc_name)

                # Allow user to download the Word document
                st.download_button(label="Download OCR Document", data=doc_name.open('rb'), file_name=doc_name.name)
            except Exception as e:
                st.error(f"Error performing OCR: {e}")

if __name__ == "__main__":
    main()

